"""
營建自動化管理系統 — 文件解析引擎
支援 Word (.docx), Excel (.xlsx), PDF 掃描檔解析
"""
import os
import re
import json
from typing import Optional


def parse_uploaded_file(filepath: str, filename: str) -> dict:
    """
    根據檔案類型自動選擇解析器
    回傳結構化的工程基本資料
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.docx':
        return parse_word(filepath)
    elif ext in ('.xlsx', '.xls'):
        return parse_excel(filepath)
    elif ext == '.pdf':
        return parse_pdf(filepath)
    elif ext in ('.jpg', '.jpeg', '.png', '.tiff', '.bmp'):
        return parse_image(filepath)
    else:
        return {"raw_text": "", "parsed_fields": {}, "error": f"不支援的檔案格式: {ext}"}


def parse_word(filepath: str) -> dict:
    """解析 Word 文件，提取工程基本資料"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "需要安裝 python-docx: pip install python-docx"}

    try:
        doc = Document(filepath)
        full_text = []
        tables_data = []

        # 提取所有段落文字
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # 提取所有表格
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append(table_rows)

        raw_text = '\n'.join(full_text)
        parsed_fields = extract_fields_from_text(raw_text)

        return {
            "raw_text": raw_text,
            "tables": tables_data,
            "parsed_fields": parsed_fields,
            "source_type": "word"
        }
    except Exception as e:
        return {"error": f"Word 解析失敗: {str(e)}"}


def parse_excel(filepath: str) -> dict:
    """解析 Excel 文件，提取工程基本資料"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"error": "需要安裝 openpyxl: pip install openpyxl"}

    try:
        wb = load_workbook(filepath, data_only=True)
        sheets_data = {}
        full_text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                rows.append(row_data)
                full_text_parts.append(' '.join([c for c in row_data if c]))
            sheets_data[sheet_name] = rows

        raw_text = '\n'.join(full_text_parts)
        parsed_fields = extract_fields_from_text(raw_text)

        # 嘗試從表格結構直接對應欄位
        for sheet_name, rows in sheets_data.items():
            for row in rows:
                if len(row) >= 2:
                    key = row[0].strip() if row[0] else ''
                    val = row[1].strip() if row[1] else ''
                    mapped = map_field_name(key)
                    if mapped and val:
                        parsed_fields[mapped] = val

        return {
            "raw_text": raw_text,
            "sheets": sheets_data,
            "parsed_fields": parsed_fields,
            "source_type": "excel"
        }
    except Exception as e:
        return {"error": f"Excel 解析失敗: {str(e)}"}


def parse_pdf(filepath: str) -> dict:
    """解析 PDF 文件（文字型 PDF 或掃描檔 OCR）"""
    raw_text = ""

    # 先嘗試文字型 PDF
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        for page in doc:
            raw_text += page.get_text() + "\n"
        doc.close()
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    raw_text += text + "\n"
        except ImportError:
            return {"error": "需要安裝 PyMuPDF 或 PyPDF2: pip install PyMuPDF 或 pip install PyPDF2"}

    # 如果文字太少，可能是掃描檔，嘗試 OCR
    if len(raw_text.strip()) < 50:
        ocr_result = _try_ocr_pdf(filepath)
        if ocr_result:
            raw_text = ocr_result

    parsed_fields = extract_fields_from_text(raw_text)

    return {
        "raw_text": raw_text,
        "parsed_fields": parsed_fields,
        "source_type": "pdf"
    }


def parse_image(filepath: str) -> dict:
    """解析圖片（OCR）"""
    raw_text = ""

    try:
        from PIL import Image
        import pytesseract
        img = Image.open(filepath)
        raw_text = pytesseract.image_to_string(img, lang='chi_tra+eng')
    except ImportError:
        return {"error": "需要安裝 Pillow 和 pytesseract: pip install Pillow pytesseract"}
    except Exception as e:
        return {"error": f"OCR 辨識失敗: {str(e)}"}

    parsed_fields = extract_fields_from_text(raw_text)

    return {
        "raw_text": raw_text,
        "parsed_fields": parsed_fields,
        "source_type": "image_ocr"
    }


def _try_ocr_pdf(filepath: str) -> Optional[str]:
    """嘗試對 PDF 進行 OCR"""
    try:
        import fitz
        from PIL import Image
        import pytesseract
        import io

        doc = fitz.open(filepath)
        full_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang='chi_tra+eng')
            full_text.append(text)
        doc.close()
        return '\n'.join(full_text)
    except Exception:
        return None


# ===== 欄位提取邏輯 =====

# 中文數字序號對應
CHINESE_NUMBERS = {
    '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
    '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
    '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15,
    '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
    '二十一': 21
}

# 欄位關鍵字對應
FIELD_PATTERNS = [
    (r'工程名稱[：:\s]*(.+)', 'project_name'),
    (r'主辦單位[：:\s]*(.+)', 'owner_agency'),
    (r'主辦機關[：:\s]*(.+)', 'owner_agency'),
    (r'執行工程處[：:\s]*(.+)', 'execution_office'),
    (r'督導工務所[：:\s]*(.+)', 'supervision_office'),
    (r'設計單位[：:\s]*(.+)', 'design_unit'),
    (r'監造單位[：:\s]*(.+)', 'supervision_unit'),
    (r'承攬廠商[：:\s]*(.+)', 'contractor'),
    (r'承攬人[：:\s]*(.+)', 'contractor'),
    (r'工程地點[：:\s]*(.+)', 'project_location'),
    (r'契約金額[：:\s]*(.+)', 'contract_amount'),
    (r'工程期限[：:\s]*(.+)', 'construction_period'),
    (r'工程範圍[：:\s]*(.+)', 'project_scope'),
    (r'保固期限[：:\s]*(.+)', 'warranty_terms'),
    (r'開工日期[：:\s]*(.+)', 'start_date'),
    (r'完工日期[：:\s]*(.+)', 'expected_end_date'),
]


def extract_fields_from_text(text: str) -> dict:
    """從文字中提取結構化欄位"""
    fields = {}

    if not text:
        return fields

    # 逐行匹配
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 移除序號前綴 如 (一)、 (1)、 1.
        cleaned = re.sub(r'^[\(（]?[\d一二三四五六七八九十]+[\)）]?[、.\s]*', '', line)

        for pattern, field_name in FIELD_PATTERNS:
            match = re.search(pattern, cleaned)
            if match:
                value = match.group(1).strip()
                if value:
                    fields[field_name] = value
                break

    # 提取數字金額
    if 'contract_amount' in fields:
        numeric = extract_numeric_amount(fields['contract_amount'])
        if numeric:
            fields['contract_amount_numeric'] = numeric

    # 提取工期天數
    if 'construction_period' in fields:
        days = extract_period_days(fields['construction_period'])
        if days:
            fields['construction_period_days'] = days

    # 提取名詞定義
    term_defs = extract_term_definitions(text)
    if term_defs:
        fields['term_definitions'] = term_defs

    return fields


def extract_numeric_amount(amount_str: str) -> Optional[float]:
    """從金額字串提取數字"""
    # 直接數字
    match = re.search(r'[\d,]+', amount_str.replace(',', ''))
    if match:
        try:
            return float(match.group().replace(',', ''))
        except ValueError:
            pass

    # 中文大寫金額轉換
    cn_digits = {
        '零': 0, '壹': 1, '貳': 2, '參': 3, '肆': 4,
        '伍': 5, '陸': 6, '柒': 7, '捌': 8, '玖': 9
    }
    cn_units = {
        '拾': 10, '佰': 100, '仟': 1000, '萬': 10000, '億': 100000000
    }

    try:
        total = 0
        current = 0
        wan_part = 0  # 萬以上的部分

        for char in amount_str:
            if char in cn_digits:
                current = cn_digits[char]
            elif char == '億':
                wan_part = (wan_part + current) * 100000000
                current = 0
            elif char == '萬':
                wan_part = wan_part + (current if current else 0) * 10000
                # 處理前面累積的
                total = (total + current) * 10000
                current = 0
                total = 0
            elif char in cn_units:
                if current == 0 and char == '拾':
                    current = 1
                total += current * cn_units[char]
                current = 0
            elif char in ('元', '整'):
                total += current
                current = 0

        total += current + wan_part
        if total > 0:
            return float(total)
    except Exception:
        pass

    return None


def extract_period_days(period_str: str) -> Optional[int]:
    """從工期字串提取天數"""
    match = re.search(r'(\d+)\s*日曆天', period_str)
    if match:
        return int(match.group(1))

    match = re.search(r'(\d+)\s*天', period_str)
    if match:
        return int(match.group(1))

    match = re.search(r'(\d+)', period_str)
    if match:
        return int(match.group(1))

    return None


def extract_term_definitions(text: str) -> list:
    """提取名詞定義列表"""
    terms = []

    # 匹配 (1)xxx-yyy 或 (1)xxx係指yyy 格式
    pattern = r'[\(（](\d+)[\)）]\s*(.+?)(?:[—\-－–]\s*|係指\s*)(.+?)(?=[\(（]\d+[\)）]|$)'
    matches = re.finditer(pattern, text, re.DOTALL)

    for match in matches:
        num = int(match.group(1))
        name = match.group(2).strip().rstrip('—-－–')
        definition = match.group(3).strip()

        # 清理定義文字
        definition = re.sub(r'\s+', ' ', definition)
        if definition.endswith('。'):
            pass
        elif len(definition) > 10:
            pass
        else:
            continue

        terms.append({
            "term_number": num,
            "term_name": name,
            "term_definition": definition
        })

    return terms


def parse_contract_xls(filepath: str) -> dict:
    """
    解析契約詳細表 Excel (.xls/.xlsx)
    回傳結構化的契約項目列表，對應 contract_items 資料表
    
    Excel 格式：
    C0=項次 | C1=項目及說明 | C2=單位 | C3=數量 | C4=單價 | C5=複價 | C6=編碼(備註)
    """
    ext = os.path.splitext(filepath)[1].lower()
    rows_data = []

    # 讀取 Excel
    if ext == '.xls':
        try:
            import xlrd
            wb = xlrd.open_workbook(filepath)
            ws = wb.sheet_by_index(0)
            for r in range(ws.nrows):
                row = []
                for c in range(ws.ncols):
                    val = ws.cell_value(r, c)
                    if isinstance(val, float) and val == int(val):
                        val = int(val)
                    row.append(val if val else '')
                rows_data.append(row)
        except ImportError:
            return {"error": "需要安裝 xlrd: pip install xlrd", "items": []}
        except Exception as e:
            return {"error": f"XLS 解析失敗: {str(e)}", "items": []}
    else:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, data_only=True)
            ws = wb.active
            for row in ws.iter_rows(values_only=True):
                row_data = []
                for cell in row:
                    val = cell if cell is not None else ''
                    if isinstance(val, float) and val == int(val):
                        val = int(val)
                    row_data.append(val)
                rows_data.append(row_data)
        except ImportError:
            return {"error": "需要安裝 openpyxl: pip install openpyxl", "items": []}
        except Exception as e:
            return {"error": f"XLSX 解析失敗: {str(e)}", "items": []}

    if not rows_data:
        return {"error": "Excel 檔案為空", "items": []}

    # 跳過標題行
    start_row = 0
    for i, row in enumerate(rows_data):
        c1 = str(row[1] if len(row) > 1 else '').strip()
        if '項' in c1 and '說明' in c1:
            start_row = i + 1
            break
        if '項目' in c1:
            start_row = i + 1
            break

    # 大類代碼辨識
    LEVEL1_CODES = {'壹', '貳', '參', '肆', '伍', '陸', '柒', '捌', '玖', '拾'}
    LEVEL2_CODES = {'甲', '乙', '丙', '丁', '戊', '己', '庚', '辛'}

    items = []
    current_l1 = ''  # 壹、貳...
    current_l2 = ''  # 甲、乙...
    current_l3 = ''  # A、B、C...
    pending_name_parts = []  # 跨行名稱暫存

    for r in range(start_row, len(rows_data)):
        row = rows_data[r]
        # 確保至少 7 欄
        while len(row) < 7:
            row.append('')

        c0 = str(row[0]).strip()
        c1 = str(row[1]).strip()
        c2 = str(row[2]).strip()
        c3 = row[3]
        c4 = row[4]
        c5 = row[5]
        c6 = str(row[6]).strip()

        # 跳過完全空行
        if not c0 and not c1 and not c2 and not c5 and not c6:
            continue

        # 處理跨行名稱（上一行名稱未完，此行只有 C1 或 C6 有值）
        if not c0 and not c2 and not c5 and (c1 or c6):
            if items and not c1.startswith('小計') and not c1.startswith('合計') and not c1.startswith('總價') and '合計' not in c1:
                # 附加到上一個項目的名稱
                if c1:
                    items[-1]['item_name'] += c1
                if c6 and items[-1].get('remark_code'):
                    items[-1]['remark_code'] += c6
                elif c6:
                    items[-1]['remark_code'] = c6
                continue

        # 判斷層級
        level = 4  # 預設為工項
        is_category = 0
        is_subtotal = 0
        parent_code = current_l3 or current_l2 or current_l1
        item_code = c0

        # Level 1: 壹、貳...
        if c0 in LEVEL1_CODES:
            level = 0
            is_category = 1
            current_l1 = c0
            current_l2 = ''
            current_l3 = ''
            parent_code = ''
            item_code = c0

        # Level 2: 甲、乙、丙、丁、戊、己、庚
        elif c0 in LEVEL2_CODES:
            level = 1
            is_category = 1
            current_l2 = c0
            current_l3 = ''
            parent_code = current_l1
            item_code = c0

        # Level 3: A, B, C... (單個大寫字母)
        elif re.match(r'^[A-Z]$', c0):
            level = 2
            is_category = 1
            current_l3 = c0
            parent_code = current_l2 or current_l1
            item_code = c0

        # Level 3 with number: Z1, Z2... (字母+數字的工項)
        elif re.match(r'^[A-Z]\d+$', c0):
            level = 3
            parent_code = current_l3 or current_l2

        # Level 4: 數字工項 1, 2, 3...
        elif re.match(r'^\d+$', c0):
            level = 3
            parent_code = current_l3 or current_l2

        # 小計/合計行
        if '小計' in c1 or '合計' in c1:
            is_subtotal = 1
            is_category = 0
            level = 2 if '小計' in c1 else 1

        if '總價' in c1 or '總計' in c1:
            is_subtotal = 1
            level = 0

        # 數值處理
        quantity = None
        unit_price = None
        total_price = None

        if c3 and c3 != '':
            try:
                quantity = float(c3)
            except (ValueError, TypeError):
                pass

        if c4 and c4 != '':
            try:
                unit_price = float(c4)
            except (ValueError, TypeError):
                pass

        if c5 and c5 != '':
            try:
                total_price = float(c5)
            except (ValueError, TypeError):
                pass

        if not c1 and not c0:
            continue

        item = {
            'row_order': len(items),
            'level': level,
            'parent_code': parent_code,
            'item_code': item_code,
            'item_number': c0,
            'item_name': c1,
            'unit': c2,
            'quantity': quantity,
            'unit_price': unit_price,
            'total_price': total_price,
            'remark_code': c6,
            'is_category': is_category,
            'is_subtotal': is_subtotal
        }
        items.append(item)

    # 統計
    work_items = [i for i in items if not i['is_category'] and not i['is_subtotal'] and i['unit']]
    categories = [i for i in items if i['is_category']]
    subtotals = [i for i in items if i['is_subtotal']]

    return {
        "items": items,
        "stats": {
            "total_rows": len(items),
            "work_items": len(work_items),
            "categories": len(categories),
            "subtotals": len(subtotals)
        }
    }


def map_field_name(key: str) -> Optional[str]:
    """將中文欄位名對應到資料庫欄位"""
    mapping = {
        '工程名稱': 'project_name',
        '主辦單位': 'owner_agency',
        '主辦機關': 'owner_agency',
        '執行工程處': 'execution_office',
        '督導工務所': 'supervision_office',
        '設計單位': 'design_unit',
        '監造單位': 'supervision_unit',
        '承攬廠商': 'contractor',
        '承攬人': 'contractor',
        '工程地點': 'project_location',
        '契約金額': 'contract_amount',
        '工程期限': 'construction_period',
        '工程範圍': 'project_scope',
        '保固期限': 'warranty_terms',
        '開工日期': 'start_date',
        '完工日期': 'expected_end_date',
    }
    return mapping.get(key.strip())
