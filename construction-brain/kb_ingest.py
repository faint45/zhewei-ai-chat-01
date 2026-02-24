# -*- coding: utf-8 -*-
"""
築未科技 Construction Brain
kb_ingest.py

功能：
  餵入大量施工法/分項計畫書/施工規範（PDF/Word/txt）
  → 文字抽取 → 分段（chunking）→ 向量索引（ChromaDB）
  → 存入本地知識庫，供 kb_query.py 查詢

用法：
    # 餵入單一文件
    python kb_ingest.py --file 橋梁施工法.pdf --category 施工法

    # 批次餵入整個資料夾
    python kb_ingest.py --folder C:\ZheweiConstruction\03_KB --category 分項計畫書

    # 查看知識庫統計
    python kb_ingest.py --stats
"""

import argparse
import hashlib
import json
import os
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path

BASE_DIR = Path(os.environ.get("ZHEWEI_BASE", r"C:\ZheweiConstruction"))
DB_PATH = BASE_DIR / "db" / "index.db"
KB_DIR = BASE_DIR / "db" / "kb_chroma"

CHUNK_SIZE = 600
CHUNK_OVERLAP = 100

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

CATEGORIES = [
    "施工法",
    "分項計畫書",
    "施工規範",
    "監造計畫",
    "設計圖說",
    "契約",
    "標準圖",
    "其他",
]


def _get_chroma_collection():
    try:
        import chromadb
        from chromadb.utils import embedding_functions

        client = chromadb.PersistentClient(path=str(KB_DIR))
        ef = embedding_functions.DefaultEmbeddingFunction()
        collection = client.get_or_create_collection(
            name="construction_kb",
            embedding_function=ef,
            metadata={"hnsw:space": "cosine"},
        )
        return collection
    except ImportError:
        print("[kb_ingest] ⚠️ chromadb 未安裝：pip install chromadb")
        return None


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return "\n".join(parts)
        except ImportError:
            print("[kb_ingest] ⚠️ pdfplumber 未安裝：pip install pdfplumber")
            return ""
        except Exception as e:
            print(f"[kb_ingest] PDF 讀取失敗：{e}")
            return ""

    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except ImportError:
            print("[kb_ingest] ⚠️ python-docx 未安裝：pip install python-docx")
            return ""
        except Exception as e:
            print(f"[kb_ingest] DOCX 讀取失敗：{e}")
            return ""

    elif suffix in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8", errors="ignore")

    return ""


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    按句子邊界切割文字，保留 overlap 字元的上下文重疊
    """
    import re
    sentences = re.split(r"(?<=[。！？\n])", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    chunks = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) <= chunk_size:
            current += sent
        else:
            if current:
                chunks.append(current)
            if len(sent) > chunk_size:
                for i in range(0, len(sent), chunk_size - overlap):
                    chunks.append(sent[i:i + chunk_size])
                current = sent[-(overlap):]
            else:
                current = (current[-overlap:] if len(current) > overlap else current) + sent

    if current:
        chunks.append(current)

    return [c for c in chunks if len(c.strip()) > 20]


def _ensure_db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kb_documents (
            id TEXT PRIMARY KEY,
            filename TEXT,
            file_path TEXT,
            file_hash TEXT UNIQUE,
            category TEXT,
            chunk_count INTEGER,
            char_count INTEGER,
            ingested_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def _is_already_ingested(file_hash: str) -> bool:
    _ensure_db()
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT id FROM kb_documents WHERE file_hash=?", (file_hash,)
    ).fetchone()
    conn.close()
    return row is not None


def ingest_file(file_path: Path, category: str = "其他", force: bool = False) -> dict:
    """
    餵入單一文件到知識庫

    Args:
        file_path: 文件路徑
        category: 文件類別（施工法/分項計畫書/施工規範/...）
        force: True = 強制重新餵入（即使已存在）

    Returns:
        結果 dict
    """
    _ensure_db()
    KB_DIR.mkdir(parents=True, exist_ok=True)
    file_path = Path(file_path)

    if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        print(f"[kb_ingest] 略過不支援的格式：{file_path.name}")
        return {"status": "skipped", "reason": "unsupported format"}

    file_hash = _file_hash(file_path)
    if not force and _is_already_ingested(file_hash):
        print(f"[kb_ingest] 已餵入，略過：{file_path.name}")
        return {"status": "duplicate", "file": file_path.name}

    print(f"[kb_ingest] 處理：{file_path.name} ({category})")
    text = _extract_text(file_path)
    if not text:
        print(f"[kb_ingest] ❌ 無法讀取文字：{file_path.name}")
        return {"status": "error", "reason": "empty text"}

    chunks = _chunk_text(text)
    print(f"[kb_ingest]   切分為 {len(chunks)} 個段落（{len(text)} 字元）")

    collection = _get_chroma_collection()
    if collection is None:
        print("[kb_ingest] ❌ ChromaDB 不可用，無法建立向量索引")
        return {"status": "error", "reason": "chromadb not available"}

    doc_id = str(uuid.uuid4())
    chunk_ids = []
    chunk_docs = []
    chunk_metas = []

    for i, chunk in enumerate(chunks):
        chunk_id = f"{doc_id}_{i:04d}"
        chunk_ids.append(chunk_id)
        chunk_docs.append(chunk)
        chunk_metas.append({
            "doc_id": doc_id,
            "filename": file_path.name,
            "category": category,
            "chunk_index": i,
            "total_chunks": len(chunks),
        })

    batch_size = 50
    for start in range(0, len(chunk_ids), batch_size):
        collection.add(
            ids=chunk_ids[start:start + batch_size],
            documents=chunk_docs[start:start + batch_size],
            metadatas=chunk_metas[start:start + batch_size],
        )

    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        """INSERT OR REPLACE INTO kb_documents
           (id, filename, file_path, file_hash, category, chunk_count, char_count, ingested_at)
           VALUES (?,?,?,?,?,?,?,?)""",
        (doc_id, file_path.name, str(file_path), file_hash,
         category, len(chunks), len(text), datetime.now().isoformat()),
    )
    conn.commit()
    conn.close()

    print(f"[kb_ingest] ✅ 完成：{file_path.name} → {len(chunks)} 段落")
    return {"status": "ok", "doc_id": doc_id, "chunks": len(chunks), "chars": len(text)}


def ingest_folder(folder_path: Path, category: str = "其他", force: bool = False) -> dict:
    """批次餵入整個資料夾"""
    folder_path = Path(folder_path)
    if not folder_path.exists():
        print(f"[kb_ingest] ❌ 資料夾不存在：{folder_path}")
        return {}

    files = [
        f for f in folder_path.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    print(f"[kb_ingest] 找到 {len(files)} 個文件，開始批次餵入...")

    results = {"ok": 0, "duplicate": 0, "error": 0, "skipped": 0}
    for f in files:
        result = ingest_file(f, category=category, force=force)
        results[result.get("status", "error")] = results.get(result.get("status", "error"), 0) + 1

    print(f"\n[kb_ingest] 批次完成：✅ {results['ok']} 個｜重複略過 {results['duplicate']} 個｜錯誤 {results['error']} 個")
    return results


def show_stats():
    """顯示知識庫統計資訊"""
    _ensure_db()
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT category, COUNT(*) as cnt, SUM(chunk_count) as chunks, SUM(char_count) as chars FROM kb_documents GROUP BY category"
    ).fetchall()
    total = conn.execute("SELECT COUNT(*), SUM(chunk_count), SUM(char_count) FROM kb_documents").fetchone()
    conn.close()

    print(f"\n{'='*55}")
    print(f"  築未科技 施工知識庫 統計")
    print(f"{'='*55}")
    print(f"{'類別':<12} {'文件數':>6} {'段落數':>8} {'字元數':>10}")
    print(f"{'-'*55}")
    for row in rows:
        print(f"{row[0]:<12} {row[1]:>6} {row[2]:>8} {row[3]:>10,}")
    print(f"{'-'*55}")
    print(f"{'合計':<12} {total[0]:>6} {total[1] or 0:>8} {total[2] or 0:>10,}")
    print(f"{'='*55}\n")

    collection = _get_chroma_collection()
    if collection:
        print(f"  向量索引段落總數：{collection.count():,}")
    print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="築未科技 — 施工知識庫餵入")
    parser.add_argument("--file", help="單一文件路徑")
    parser.add_argument("--folder", help="批次餵入整個資料夾")
    parser.add_argument("--category", default="其他",
                        choices=CATEGORIES, help="文件類別")
    parser.add_argument("--force", action="store_true", help="強制重新餵入（覆蓋已有）")
    parser.add_argument("--stats", action="store_true", help="顯示知識庫統計")
    args = parser.parse_args()

    if args.stats:
        show_stats()
    elif args.file:
        ingest_file(Path(args.file), category=args.category, force=args.force)
    elif args.folder:
        ingest_folder(Path(args.folder), category=args.category, force=args.force)
    else:
        parser.print_help()
