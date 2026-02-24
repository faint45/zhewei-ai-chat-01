# -*- coding: utf-8 -*-
"""
築未科技 Construction Brain
schedule_extractor.py

功能：
  上傳契約/監造計畫/設計圖說（PDF/Word/txt）
  → AI 自動辨識工項、工期、邏輯關係
  → 輸出 schedule.json（供 schedule_engine.py 計算 CPM）
  → 同步寫入 SQLite

用法：
    python schedule_extractor.py --file contract.pdf --project_id PRJ-001
    python schedule_extractor.py --file spec.docx --project_id PRJ-001 --append
"""

import argparse
import json
import os
import re
import sqlite3
import uuid
from datetime import datetime, date
from pathlib import Path

import httpx

BASE_DIR = Path(os.environ.get("ZHEWEI_BASE", r"C:\ZheweiConstruction"))
DB_PATH = BASE_DIR / "db" / "index.db"
OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "zhewei-brain")
OLLAMA_TIMEOUT = float(os.environ.get("OLLAMA_TIMEOUT", "180"))

EXTRACT_SYSTEM_PROMPT = """你是「築未科技」的工程進度規劃引擎。

【任務】
從輸入的工程文件（契約、監造計畫、設計圖說、施工規範）中，
抽取出完整的「工項清單」，並推斷合理的施工邏輯關係。

【你必須輸出純 JSON，不得附加任何說明文字】格式如下：

{
  "project_name": "工程名稱（從文件中讀取，找不到填null）",
  "contract_period_days": 工期天數或null,
  "start_date_hint": "文件提到的開工日或null",
  "work_items": [
    {
      "id": "WBS編號，如1/1.1/1.1.1，無則自動編A001起",
      "name": "工項名稱",
      "duration_days": 估計工期天數（整數，合理範圍1-180）,
      "predecessors": ["前置工項id清單，無前置填[]"],
      "trade": "負責工種/分包類別（土木/鋼筋/模板/混凝土/機電/道路/橋梁/管線/其他）",
      "unit": "數量單位（m/m²/m³/噸/式/組/項）",
      "quantity": 數量或null,
      "weight_pct": 工程費占比百分比或null,
      "notes": "備註（特殊施工條件/規範要求）"
    }
  ],
  "critical_constraints": ["重要施工限制條件，如：主橋段需交通管制/雨季停工等"],
  "parse_confidence": "high|medium|low（對文件內容的解析信心程度）"
}

【抽取規則】
1. 工項粒度：每個「可獨立施工、有明確完工標準」的作業為一個工項
   - 太粗（整個橋梁）→ 細分到分部（基礎/墩柱/上部結構）
   - 太細（每根螺栓）→ 合併為一個工作包
2. 工期估算原則（無明確工期時）：
   - 開挖整地：數量/300m³ 每天，最少2天
   - 鋼筋綁紮：數量/5噸 每天，最少3天
   - 混凝土澆置：數量/50m³ 每天，最少1天（含養護7天）
   - 模板：與混凝土工項相同工期
   - 道路鋪面：每公里3-5天
   - 管線：每公里5-10天
3. 邏輯關係原則（FS=完成後才能開始）：
   - 基礎完成 → 墩柱；墩柱完成 → 帽梁；帽梁完成 → 上部結構
   - 土方開挖 → 地下結構；地下結構 → 回填
   - 路基整備 → 底層鋪設 → 面層鋪設
4. weight_pct：從契約金額分項或工程費用表讀取；無則填 null
5. 禁止輸出 JSON 以外任何文字
6. 工項總數建議 10-50 項（太少就細分，太多就合併）"""


def _extract_text_from_pdf(path: Path) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except ImportError:
        print("[schedule_extractor] ⚠️ pdfplumber 未安裝：pip install pdfplumber")
        return ""
    except Exception as e:
        print(f"[schedule_extractor] PDF 讀取失敗：{e}")
        return ""


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        print("[schedule_extractor] ⚠️ python-docx 未安裝：pip install python-docx")
        return ""
    except Exception as e:
        print(f"[schedule_extractor] DOCX 讀取失敗：{e}")
        return ""


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        text = _extract_text_from_docx(file_path)
    elif suffix in (".txt", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        print(f"[schedule_extractor] 不支援的格式：{suffix}")
        return ""
    print(f"[schedule_extractor] 讀取 {file_path.name}：{len(text)} 字元")
    return text


def _call_ollama_schedule(text: str) -> str:
    max_chars = 12000
    if len(text) > max_chars:
        print(f"[schedule_extractor] 文件過長，截取前 {max_chars} 字元")
        text = text[:max_chars]

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": EXTRACT_SYSTEM_PROMPT},
            {"role": "user", "content": f"請從以下工程文件抽取工項清單：\n\n{text}"},
        ],
        "stream": False,
        "options": {"temperature": 0.1, "num_ctx": 16384},
    }
    with httpx.Client(timeout=OLLAMA_TIMEOUT) as client:
        r = client.post(f"{OLLAMA_BASE_URL}/api/chat", json=payload)
        r.raise_for_status()
        return r.json()["message"]["content"].strip()


def _parse_json_response(raw: str) -> dict:
    raw = raw.strip()
    m = re.search(r"\{[\s\S]*\}", raw)
    if m:
        return json.loads(m.group(0))
    raise ValueError(f"找不到 JSON：{raw[:200]}")


def _auto_assign_ids(work_items: list) -> list:
    for i, item in enumerate(work_items):
        if not item.get("id"):
            item["id"] = f"A{i+1:03d}"
    return work_items


def _validate_predecessors(work_items: list) -> list:
    valid_ids = {item["id"] for item in work_items}
    for item in work_items:
        item["predecessors"] = [
            p for p in item.get("predecessors", []) if p in valid_ids
        ]
    return work_items


def _ensure_db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS schedules (
            id TEXT PRIMARY KEY,
            project_id TEXT,
            source_file TEXT,
            schedule_json TEXT,
            parse_confidence TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def extract_schedule(
    file_path: Path,
    project_id: str = "default",
    append: bool = False,
) -> dict:
    """
    主要入口：從文件抽取工項清單，寫入 schedule.json 與 SQLite

    Args:
        file_path: 上傳的文件路徑（PDF/Word/txt）
        project_id: 專案代碼
        append: True = 合併到現有 schedule，False = 覆蓋

    Returns:
        schedule dict
    """
    _ensure_db()
    file_path = Path(file_path)

    text = extract_text(file_path)
    if not text:
        print("[schedule_extractor] ❌ 文件內容為空，無法抽取")
        return {}

    print(f"[schedule_extractor] 呼叫 AI 抽取工項（可能需要 30-120 秒）...")
    try:
        raw = _call_ollama_schedule(text)
        schedule = _parse_json_response(raw)
    except Exception as e:
        print(f"[schedule_extractor] ❌ AI 抽取失敗：{e}")
        return {"error": str(e), "raw": raw if "raw" in dir() else ""}

    work_items = schedule.get("work_items", [])
    work_items = _auto_assign_ids(work_items)
    work_items = _validate_predecessors(work_items)
    schedule["work_items"] = work_items

    out_dir = BASE_DIR / "projects" / project_id / "02_Output" / "Schedule"
    out_dir.mkdir(parents=True, exist_ok=True)
    schedule_path = out_dir / "schedule.json"

    if append and schedule_path.exists():
        existing = json.loads(schedule_path.read_text(encoding="utf-8"))
        existing_items = existing.get("work_items", [])
        existing_ids = {item["id"] for item in existing_items}
        new_items = [item for item in work_items if item["id"] not in existing_ids]
        existing_items.extend(new_items)
        schedule["work_items"] = existing_items
        print(f"[schedule_extractor] 合併模式：新增 {len(new_items)} 個工項")

    schedule_path.write_text(json.dumps(schedule, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[schedule_extractor] ✅ schedule.json → {schedule_path}")
    print(f"[schedule_extractor]    工項數：{len(schedule['work_items'])}　信心度：{schedule.get('parse_confidence','?')}")

    schedule_id = str(uuid.uuid4())
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        """INSERT INTO schedules (id, project_id, source_file, schedule_json, parse_confidence, created_at)
           VALUES (?,?,?,?,?,?)""",
        (schedule_id, project_id, file_path.name,
         json.dumps(schedule, ensure_ascii=False),
         schedule.get("parse_confidence", ""),
         datetime.now().isoformat()),
    )
    conn.commit()
    conn.close()

    csv_path = out_dir / "schedule.csv"
    _write_schedule_csv(schedule["work_items"], csv_path)

    return schedule


def _write_schedule_csv(work_items: list, csv_path: Path):
    import csv
    fieldnames = ["id", "name", "duration_days", "predecessors", "trade",
                  "unit", "quantity", "weight_pct", "notes"]
    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for item in work_items:
            row = dict(item)
            row["predecessors"] = ",".join(item.get("predecessors", []))
            writer.writerow(row)
    print(f"[schedule_extractor] ✅ schedule.csv → {csv_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="築未科技 — AI 工程進度工項抽取")
    parser.add_argument("--file", required=True, help="文件路徑（PDF/Word/txt）")
    parser.add_argument("--project_id", default="test_project_001")
    parser.add_argument("--append", action="store_true", help="合併到現有工項（不覆蓋）")
    args = parser.parse_args()

    result = extract_schedule(
        file_path=Path(args.file),
        project_id=args.project_id,
        append=args.append,
    )
    if result:
        print(f"\n--- 抽取摘要 ---")
        print(f"工程名稱：{result.get('project_name','')}")
        print(f"核定工期：{result.get('contract_period_days','')} 天")
        print(f"工項數量：{len(result.get('work_items',[]))} 項")
        print(f"施工限制：{result.get('critical_constraints',[])}")
